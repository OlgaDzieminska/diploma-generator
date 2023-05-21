import argparse
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, Inches

parser = argparse.ArgumentParser()
parser.add_argument("form_data")
parser.add_argument("document_name")
args = parser.parse_args()

inp = args.form_data
document_name = args.document_name
form_data = json.loads(inp)

font_name = 'Times New Roman'
font_size = 12


def createDocument(font_name, font_size, form_data, document_name):
    document = Document()
    __appendTitleOfDocument(document)
    __appendPupilDataConclusionAndRecommendation(document, form_data)
    __appendDate(document, form_data)
    document.save(document_name)
    print("Created document:" + document_name)


def set_Font(paragraph, _font_size, font_name):
    font = paragraph.runs[0].font
    font.name = font_name
    font.size = Pt(_font_size)


def __appendTitleOfDocument(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run('OCENA EFEKTYWNOŚCI UDZIELANEJ POMOCY\n'
                            'PSYCHOLOGICZNO-PEDAGOGICZNEJ\n'
                            'w roku szkolnym 2022/2023')
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font = run.font
    font.bold = True
    set_Font(paragraph, font_size, font_name)


def __appendPupilDataConclusionAndRecommendation(document, _form_data):
    conclusion = " ".join(form_data['conclusionParts'])
    recommendation = " ".join(form_data['recommendationParts'])
    class_number = _form_data['pupilDTO']['classNumber']
    pupil_name = _form_data['pupilDTO']['name']
    teacher_name = _form_data['teacher']
    lecture = _form_data['lecture']
    elements = ['Nazwisko i imię ucznia: ' + pupil_name + '\nKlasa:' + class_number,
                'Rodzaj zajęć: ' + lecture,
                'Nauczyciel prowadzący zajęcia: ' + teacher_name,
                'Wnioski ( co się poprawiło lub uległo pogorszeniu? ):\n' + conclusion,
                'Zalecenia dotyczące dalszych działań mających na celu poprawę funkcjonowania ucznia:\n' + recommendation]
    for i in elements:
        paragraph = document.add_paragraph(i, style='List Number')
        set_Font(paragraph, font_size, font_name)


def __appendDate(document, _form_data):
    _date = _form_data['date']
    date = document.add_paragraph()
    run = date.add_run('Data:\n' + _date)
    date.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date.paragraph_format.left_indent = Inches(0.5)
    set_Font(date, font_size, font_name)


createDocument(font_name, font_size, form_data, document_name)
